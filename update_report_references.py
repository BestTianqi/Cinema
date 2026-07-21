from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX_PATH = "report.docx"
OUTPUT_PATH = "report_updated.docx"

references = [
    "课程大作业说明：《SmartCinema 智能影院选座系统（作业说明）》。用途：确定项目功能、交互、视觉及 AI 协同开发要求。",
    "MDN Web Docs. Canvas API. https://developer.mozilla.org/zh-CN/docs/Web/API/Canvas_API （访问日期：2026-07-18）。用途：座位图与热度图的 Canvas 绘制。",
    "MDN Web Docs. Web Storage API. https://developer.mozilla.org/zh-CN/docs/Web/API/Web_Storage_API （访问日期：2026-07-18）。用途：用户、订单、售座与无障碍偏好的本地存储。",
    "MDN Web Docs. Pointer events. https://developer.mozilla.org/zh-CN/docs/Web/API/Pointer_events （访问日期：2026-07-18）。用途：点击、拖拽框选和画布平移交互。",
    "MDN Web Docs. Web Speech API. https://developer.mozilla.org/zh-CN/docs/Web/API/Web_Speech_API （访问日期：2026-07-18）。用途：无障碍语音提示。",
    "Krug, Steve. Don't Make Me Think: A Common Sense Approach to Web Usability, 3rd Edition. New Riders, 2014. 用途：简化选座步骤、降低用户决策成本。",
    "猫眼电影：《猫眼平台用户服务协议》。https://m.maoyan.com/terms/terms （访问日期：2026-07-18）。用途：参考综合在线票务、在线选座购票与评价服务的产品结构。",
    "淘票票：《淘票票—电影、演出一站式购票平台》官方应用介绍。https://apps.apple.com/cn/app/id566813949 （访问日期：2026-07-18）。用途：参考主流购票平台的在线选座流程和信息组织。",
    "中华人民共和国住房和城乡建设部：《电影院建筑设计规范》JGJ 58-2008，第 4.2 节“观众厅”。https://ebook.chinabuilding.com.cn/zbooklib/bookpdf/probation?SiteID=1&bookID=57852 （访问日期：2026-07-18）。用途：参考观众厅座位排列、排距、错位与通道设计。",
    "央视新闻客户端：《看电影什么位置最好？（不是正中间）》，中国新闻网转载，2025-02-04。https://www.chinanews.com.cn/sh/2025/02-04/10363005.shtml （访问日期：2026-07-18）。用途：确定不同规模影厅的排数范围及中后排优先策略。",
    "数字北京科学中心、上海科协、科普中国：《看电影中间位置最好？大部分的人都选错了……》，澎湃新闻，2024-03-04。https://www.thepaper.cn/newsDetail_forward_26545244 （访问日期：2026-07-18）。用途：参考过道、出入便利性和最佳观影区判断。",
    "Apple 官方网站。https://www.apple.com.cn/ （访问日期：2026-07-18）。用途：参考简洁层级、留白与高可读性界面。",
    "Tesla 官方网站。https://www.tesla.cn/ （访问日期：2026-07-18）。用途：参考科技感、深色视觉与突出主操作的设计。",
    "OpenAI 官方网站。https://openai.com/ （访问日期：2026-07-18）。用途：参考信息层级、卡片布局与克制的视觉语言。",
    "Netflix 官方网站。https://www.netflix.com/cn/ （访问日期：2026-07-18）。用途：参考影视产品的深色背景、内容聚焦与高对比交互。",
]


def set_east_asia_font(run, font_name="宋体", size=10.5):
    run.font.name = font_name
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


doc = Document(DOCX_PATH)
heading_index = next(
    i
    for i, paragraph in enumerate(doc.paragraphs)
    if paragraph.text.strip().startswith("十、参考资料")
)
heading = doc.paragraphs[heading_index]

for paragraph in list(doc.paragraphs[heading_index + 1 :]):
    paragraph._element.getparent().remove(paragraph._element)

for item in references:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(item)
    set_east_asia_font(run)

doc.save(OUTPUT_PATH)
print(f"Updated {OUTPUT_PATH} with {len(references)} references.")
