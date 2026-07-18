from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, PageBreak

DOCX = r"大作业\report.docx"
PDF = r"大作业\report.pdf"

SECTIONS = [
("一、项目概述", [
"SmartCinema 是一个无后端依赖的智能影院选座系统。项目从产品经理、UX 设计师和前端工程师三个角度出发，目标是让不同年龄、不同同行关系的观众在三步以内完成选座与购票，降低用户比较座位和理解规则的成本。",
"系统使用 HTML5、CSS3 和原生 JavaScript 编写，全部功能集中在 index.html 中，可直接使用最新版 Chrome、Edge 或 Safari 打开。用户、订单、售座和无障碍偏好均保存在 LocalStorage。"
]),
("二、用户分析", [
"大学生和普通成年人：希望快速决策、预算清晰，适合默认的综合最佳视听推荐。",
"情侣：偏好中间区域连续双座，系统优先搜索第 6 排附近的中心连座。",
"家庭：需要多人连续座位并照顾儿童或老人，系统优先中后排，遇到特殊年龄成员时自动收紧可选排数。",
"老年人：避免最后三排；界面提供大字体、高对比度和语音提示，减少阅读与操作负担。",
"少年：15 岁以下避免前三排，降低近距离观影造成的不适。",
"团体观众：5 至 20 人必须同排连续，系统会在符合全体年龄限制的排中寻找连续空位。"
]),
("三、产品设计与三步操作流程", [
"步骤 1：登录或注册后选择放映厅、票型并填写观众姓名和年龄，点击“智能推荐座位”。",
"步骤 2：查看推荐理由和体验评分；若不满意，可点击座位、按 Ctrl 多选或使用拖拽框选调整。",
"步骤 3：确认座位与金额，选择“预订座位”或“立即购票”。完成后可在订单中心取消预订或退票。",
"这一流程把规则判断、座位比较和价格计算交给系统，用户主要负责表达需求、确认结果和提交订单。"
]),
("四、视觉与交互设计", [
"色彩方案：深黑蓝背景表现影院环境和科技感；青色用于主要操作；绿色为空座、黄色为已选、红色为已售、紫色为智能推荐。热度图严格使用红、黄、蓝表示热门、一般和冷门。",
"布局设计：PC 端左侧为大面积座位画布，右侧为三步操作栏；平板变为上下混合布局；手机端改为单列，导航可横向滚动，页面无横向溢出。",
"字体选择：使用 Inter、微软雅黑和 Arial 的系统字体组合，无外部字体依赖。标题强调层级，辅助信息降低亮度，保证深色界面中的可读性。",
"交互风格：按钮状态、推荐座位、评分圆环和订单反馈均即时更新；重要操作提供文字提示，可选语音播报。"
]),
("五、功能模块实现", [
"登录注册与管理员：首次载入内置管理员 admin / admin123；普通用户注册后获得会员资格。未登录用户不能使用推荐、订单和后台功能。管理员可查看用户数、订单数、售座数和订单明细。",
"Canvas 弧形座位图：小厅 10×10、中厅 20×10、大厅 30×10。每个座位包含排号和座号，横向位置线性分布，纵向叠加基于离中心距离平方的弧形偏移。Canvas 同时承担绘制和命中检测，未使用任何第三方图表库。",
"智能推荐：先根据少年不能坐前三排、老年人不能坐最后三排确定有效排区间，再按票型人数遍历同排连续空位。候选组合按照与目标排、横向中心的距离计算分数并排序，选择最高分组合。",
"手动选座：普通点击用于单选；Ctrl 或 Command 点击保留原选择实现多选；拖拽框选模式根据矩形范围批量加入座位；支持画布平移和滚轮缩放。",
"影院热度地图：使用第二个 Canvas 根据星期、排数、列位置和确定性波动算法计算热度。周末提高中心区权重，用户可以切换星期和三种放映厅。",
"观影体验评分：系统分别计算中心视角、与第 6 排的距离和周围已售座位密度，按 40%、35%、25% 加权；再与用户 1-5 星评分合成，输出极佳、优秀或一般。",
"无障碍模式：支持大字体、高对比度、色盲友好配色和 Web Speech API 语音提示，偏好会保存在浏览器中。",
"订单中心：支持预订、购票、取消预订和退票。购票后座位进入已售状态；退票后释放座位。订单和售座状态均持久化到 LocalStorage。"
]),
("六、关键算法与数据结构", [
"放映厅配置 halls 保存名称、列数和固定 10 排；座位使用 row、col、id、sold、x、y 和 r 描述。",
"用户表 sc_users 按用户名保存密码、角色和会员状态；sc_session 保存当前登录用户；sc_orders 保存订单数组；sc_sold 按厅保存已售座位编号。",
"推荐算法的时间复杂度约为 O(排数×列数×票数)，最大厅仅 300 座，浏览器中可以即时完成。",
"评分和推荐都基于实时座位状态重新计算，推荐不满意时手动调整会立即更新座位文本、金额和体验评分。"
]),
("七、AI 协同开发说明", [
"产品设计过程：AI 用于拆解题目中的角色、模块和评分点，生成首版功能清单；学生根据三步完成目标将座位图设为首屏核心，压缩不必要页面跳转。",
"用户分析过程：AI 辅助归纳大学生、情侣、家庭、老人、少年和团体的差异；学生将年龄限制转换为明确的有效排区间，并决定团体不可拆排。",
"AI 生成内容：首版单文件结构、Canvas 绘制方法、LocalStorage 数据结构和响应式 CSS 草案由 AI 辅助生成。",
"学生修改内容：学生决定采用 0.33 GB 等基础实验配置与本项目无关；本项目中重点修改了推荐优先级、三步流程、座位色彩、深色界面密度、管理员入口和本地订单闭环，并通过实际浏览器操作检查结果。",
"用户体验优化：把推荐理由、系统评分、金额和选中座位放在同一操作栏；为手机端取消双栏；增加拖拽框选、语音提示、画布缩放和平移。",
"设计决策：不引入框架、网络接口和第三方图表库，确保压缩包解压后直接打开即可评分；用确定性算法代替随机热度，保证同一星期的结果稳定可解释。"
]),
("八、测试结果", [
"已验证普通用户注册并自动登录；会员身份显示正确。",
"已验证情侣票推荐第 6 排中心连续双座，推荐理由、评分和金额同步更新。",
"已验证购票后订单中心生成记录，座位变为已售；退票可释放座位。",
"已验证热度地图可切换一周七天和三个放映厅，Canvas 尺寸有效。",
"已验证管理员账号登录后出现后台入口，可查看统计和订单明细。",
"已验证浏览器控制台无 JavaScript 错误；390px 手机宽度下没有横向溢出。"
]),
("九、使用说明", [
"解压提交包后，直接使用最新版浏览器打开 index.html。",
"普通用户点击“注册”，输入用户名和至少 4 位密码即可进入；管理员账号为 admin，密码为 admin123。",
"若要恢复初始状态，可用管理员登录后在后台点击“重置演示数据”；也可以清除浏览器对该文件的站点数据。",
"系统所有记录仅保存在当前浏览器，不会上传到网络。"
]),
("十、参考资料", [
"课程大作业说明：《SmartCinema 智能影院选座系统》。",
"MDN Web Docs：Canvas API、Web Storage API、Pointer Events、Web Speech API。",
"Steve Krug，《Don't Make Me Think》。",
"界面风格参考 Apple、Tesla、OpenAI、Netflix、猫眼电影和淘票票公开页面；未复制其代码或素材。",
"开发过程中使用 OpenAI Codex 辅助需求分析、代码生成、文档整理和浏览器测试，最终功能取舍与修改由学生完成。"
])
]

def set_font(run, size=11, bold=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor(0,0,0)

doc=Document(); sec=doc.sections[0]
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(8); normal.paragraph_format.line_spacing=1.15
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3); set_font(p.add_run("SmartCinema 智能影院选座系统"),20,True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(12); set_font(p.add_run("大作业设计与实现说明"),12)
for k,v in [("姓名","周皓洋"),("学号","2024012535"),("联系邮箱","zhou-hy24@mails.tsinghua.edu.cn"),("入口文件","index.html，可直接使用最新版浏览器打开")]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); set_font(p.add_run(k+"："),11,True); set_font(p.add_run(v))
for title,items in SECTIONS:
    p=doc.add_heading(title,level=1); p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6); p.paragraph_format.keep_with_next=True
    for r in p.runs:set_font(r,16,True)
    for text in items:
        p=doc.add_paragraph(style="List Bullet" if len(items)>2 else None); p.paragraph_format.space_after=Pt(5); p.paragraph_format.line_spacing=1.15; set_font(p.add_run(text))
doc.core_properties.title="SmartCinema 智能影院选座系统大作业说明"; doc.core_properties.author="周皓洋"; doc.save(DOCX)

pdfmetrics.registerFont(TTFont("YaHei",r"C:\Windows\Fonts\msyh.ttc",subfontIndex=0))
pdfmetrics.registerFont(TTFont("YaHeiBold",r"C:\Windows\Fonts\msyhbd.ttc",subfontIndex=0))
B=ParagraphStyle("B",fontName="YaHei",fontSize=9.6,leading=14.2,spaceAfter=4,textColor=colors.black)
H=ParagraphStyle("H",parent=B,fontName="YaHeiBold",fontSize=14.5,leading=20,spaceBefore=10,spaceAfter=6)
T=ParagraphStyle("T",parent=B,fontName="YaHeiBold",fontSize=20,leading=28,alignment=TA_CENTER,spaceAfter=4)
S=ParagraphStyle("S",parent=B,fontSize=12,leading=18,alignment=TA_CENTER,spaceAfter=14)
story=[Paragraph("SmartCinema 智能影院选座系统",T),Paragraph("大作业设计与实现说明",S)]
for k,v in [("姓名","周皓洋"),("学号","2024012535"),("联系邮箱","zhou-hy24@mails.tsinghua.edu.cn"),("入口文件","index.html，可直接使用最新版浏览器打开")]: story.append(Paragraph(f"<b>{k}：</b>{v}",B))
for title,items in SECTIONS:
    story.append(Paragraph(title,H))
    story.append(ListFlowable([ListItem(Paragraph(x,B),leftIndent=10) for x in items],bulletType="bullet",leftIndent=18,bulletFontName="YaHei",bulletFontSize=7,spaceAfter=4))
def footer(c,d):
    c.saveState(); c.setFont("YaHei",8); c.setFillColor(colors.HexColor("#555555")); c.drawString(inch,0.48*inch,"SmartCinema 智能影院选座系统"); c.drawRightString(7.5*inch,0.48*inch,f"周皓洋 · 2024012535 · 第 {d.page} 页"); c.restoreState()
SimpleDocTemplate(PDF,pagesize=letter,leftMargin=inch,rightMargin=inch,topMargin=.72*inch,bottomMargin=.72*inch,title="SmartCinema大作业说明",author="周皓洋").build(story,onFirstPage=footer,onLaterPages=footer)
print(DOCX,PDF)
